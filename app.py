import os
import io
import json
import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, date, timedelta
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Intuit brand palette ──────────────────────────────────────────────────────
BLUE    = "#0077C5"
DARK    = "#1A1A2E"
ORANGE  = "#FF6900"
GREEN   = "#2CA01C"
RED     = "#D0021B"
TEAL    = "#007D8A"
WHITE   = "#FFFFFF"
OFFWHITE= "#F8FAFB"
BORDER  = "#D1D9E0"

# ── Session persistence ───────────────────────────────────────────────────────
CACHE_DIR = os.path.expanduser("~/.hiring_dashboard")
os.makedirs(CACHE_DIR, exist_ok=True)

# ── Standard file paths ───────────────────────────────────────────────────────
# Files must be placed in the same folder as app.py (alongside it).
APP_DIR    = os.path.dirname(os.path.abspath(__file__))
STD_REQS   = os.path.join(APP_DIR, "open_reqs.xlsx")
STD_HIRES  = os.path.join(APP_DIR, "expected_hires.xlsx")
STD_ACTUAL = os.path.join(APP_DIR, "actual_hires.xlsx")

def load_from_folder(path: str, cache_key: str):
    """Load a standard-named file from the app folder and sync it to cache."""
    if not os.path.exists(path):
        return None
    with open(path, "rb") as f:
        data = f.read()
    # Sync to cache so timestamp banner stays accurate
    save_to_cache(cache_key, data)
    return data

def _meta_path():
    return os.path.join(CACHE_DIR, "meta.json")

def load_meta() -> dict:
    p = _meta_path()
    if os.path.exists(p):
        with open(p) as f:
            return json.load(f)
    return {}

def save_to_cache(key: str, data: bytes):
    with open(os.path.join(CACHE_DIR, f"{key}.xlsx"), "wb") as f:
        f.write(data)
    meta = load_meta()
    meta[key] = datetime.now().strftime("%d %b %Y, %H:%M")
    with open(_meta_path(), "w") as f:
        json.dump(meta, f)

def load_from_cache(key: str):
    p = os.path.join(CACHE_DIR, f"{key}.xlsx")
    return open(p, "rb").read() if os.path.exists(p) else None

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Intuit Hiring Dashboard",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown(f"""
<style>
  /* ── Palette:  warm light background, always dark text, blue accents only ── */

  /* Page & main content */
  .stApp, .main, .block-container {{
    background-color: #F7F9FB !important;
  }}

  /* Sidebar */
  section[data-testid="stSidebar"],
  section[data-testid="stSidebar"] > div {{
    background-color: #EDF2F7 !important;
    border-right: 2px solid #C9D6E3 !important;
  }}

  /* ── Dark text everywhere — no exceptions ── */
  html, body, p, span, div, label, li, a, small,
  strong, em, h1, h2, h3, h4, h5,
  .stMarkdown, .stMarkdown *,
  [data-testid="stMarkdownContainer"],
  [data-testid="stMarkdownContainer"] *,
  section[data-testid="stSidebar"] *,
  .stSelectbox *, .stMultiSelect *,
  .stRadio *, .stFileUploader *,
  .stCaption, [data-testid="stCaptionContainer"] * {{
    color: #1C2B3A !important;
    font-family: Arial, sans-serif !important;
  }}

  /* ── Tabs: underline style — no color inversion ── */
  .stTabs [data-baseweb="tab-list"] {{
    background-color: #EDF2F7 !important;
    border-bottom: 3px solid {BLUE} !important;
    gap: 2px !important;
    padding-bottom: 0 !important;
  }}
  .stTabs [data-baseweb="tab"],
  .stTabs [data-baseweb="tab"] * {{
    font-size: 0.9rem !important;
    font-weight: 700 !important;
    color: #3A4A5C !important;
    background-color: #EDF2F7 !important;
    padding: 10px 28px !important;
    border-radius: 6px 6px 0 0 !important;
    border: 1px solid #C9D6E3 !important;
    border-bottom: none !important;
  }}
  /* Active tab: white card popping out of the blue bottom border */
  .stTabs [aria-selected="true"],
  .stTabs [aria-selected="true"] * {{
    color: #0077C5 !important;
    background-color: #FFFFFF !important;
    border-color: #C9D6E3 !important;
    font-weight: 800 !important;
  }}
  /* Tab content panel */
  .stTabs [data-baseweb="tab-panel"] {{
    background-color: #FFFFFF !important;
    border: 1px solid #C9D6E3 !important;
    border-top: none !important;
    border-radius: 0 0 8px 8px !important;
    padding: 24px !important;
  }}

  /* ── Metric cards ── */
  [data-testid="metric-container"] {{
    background: #FFFFFF !important;
    border: 1px solid #C9D6E3 !important;
    border-left: 5px solid {BLUE} !important;
    border-radius: 6px !important;
    padding: 14px 18px !important;
    box-shadow: 0 1px 4px rgba(0,60,120,.07) !important;
  }}
  [data-testid="metric-container"] [data-testid="stMetricLabel"] *  {{
    color: #3A4A5C !important;
    font-size: 0.78rem !important;
    font-weight: 700 !important;
    text-transform: uppercase !important;
    letter-spacing: .05em !important;
  }}
  [data-testid="metric-container"] [data-testid="stMetricValue"],
  [data-testid="metric-container"] [data-testid="stMetricValue"] * {{
    color: #1C2B3A !important;
    font-size: 1.6rem !important;
    font-weight: 800 !important;
  }}

  /* ── Headings ── */
  h1 {{ color: {BLUE} !important; font-size: 1.7rem !important; font-weight: 800 !important; }}
  h2, h3 {{ color: #1C2B3A !important; font-weight: 700 !important; }}

  /* ── Filter card ── */
  .filter-card {{
    background: #EDF2F7 !important;
    border: 1px solid #C9D6E3 !important;
    border-left: 5px solid {BLUE} !important;
    border-radius: 6px !important;
    padding: 16px 20px !important;
    margin-bottom: 20px !important;
  }}
  .filter-card * {{ color: #1C2B3A !important; }}

  /* ── Timestamp banner ── */
  .ts-banner {{
    background: #E8F0F8 !important;
    border: 1px solid #A8BDD0 !important;
    border-left: 5px solid {BLUE} !important;
    border-radius: 6px !important;
    padding: 10px 16px !important;
    font-size: 0.85rem !important;
    color: #1C2B3A !important;
    margin-bottom: 18px !important;
  }}
  .ts-banner * {{ color: #1C2B3A !important; }}
  .ts-banner strong {{ font-weight: 800 !important; }}

  /* ── Section label ── */
  .section-label {{
    font-size: 0.75rem !important;
    font-weight: 800 !important;
    text-transform: uppercase !important;
    letter-spacing: .07em !important;
    color: {BLUE} !important;
    margin-bottom: 4px !important;
  }}

  /* ── Download button ── */
  .stDownloadButton > button {{
    background-color: {BLUE} !important;
    color: #FFFFFF !important;
    border-radius: 4px !important;
    border: none !important;
    font-weight: 700 !important;
  }}
  .stDownloadButton > button * {{ color: #FFFFFF !important; }}
  .stDownloadButton > button:hover {{ background-color: #005EA6 !important; }}

  /* ── Multiselect tags: dark bg, white text (small pill — readable) ── */
  .stMultiSelect [data-baseweb="tag"] {{
    background-color: {BLUE} !important;
  }}
  .stMultiSelect [data-baseweb="tag"] span,
  .stMultiSelect [data-baseweb="tag"] svg {{
    color: #FFFFFF !important;
    fill: #FFFFFF !important;
  }}

  /* ── Dataframe ── */
  [data-testid="stDataFrame"] {{
    border: 1px solid #C9D6E3 !important;
    border-radius: 6px !important;
  }}

  /* ── Divider ── */
  hr {{ border-color: #C9D6E3 !important; margin: 20px 0 !important; }}

  /* ── Dropdown / select / multiselect overlays ── */

  /* The input box itself */
  [data-baseweb="select"] > div,
  [data-baseweb="select"] input {{
    background-color: #FFFFFF !important;
    color: #1C2B3A !important;
    border-color: #A8BDD0 !important;
  }}
  [data-baseweb="select"] > div *,
  [data-baseweb="select"] input * {{
    color: #1C2B3A !important;
  }}

  /* Placeholder text */
  [data-baseweb="select"] input::placeholder {{
    color: #6B7F96 !important;
  }}

  /* The floating dropdown popover */
  [data-baseweb="popover"],
  [data-baseweb="popover"] > div,
  ul[data-baseweb="menu"],
  [data-baseweb="menu"] {{
    background-color: #FFFFFF !important;
    border: 1px solid #A8BDD0 !important;
    border-radius: 6px !important;
    box-shadow: 0 4px 12px rgba(0,60,120,.12) !important;
  }}

  /* Each dropdown option */
  [role="option"],
  [data-baseweb="menu"] li,
  ul[data-baseweb="menu"] li {{
    background-color: #FFFFFF !important;
    color: #1C2B3A !important;
    font-family: Arial, sans-serif !important;
    font-size: 0.9rem !important;
  }}
  [role="option"] *,
  [data-baseweb="menu"] li * {{
    color: #1C2B3A !important;
  }}

  /* Hover state on option */
  [role="option"]:hover,
  [data-baseweb="menu"] li:hover {{
    background-color: #E8F0F8 !important;
    color: #0077C5 !important;
    cursor: pointer !important;
  }}
  [role="option"]:hover *,
  [data-baseweb="menu"] li:hover * {{
    color: #0077C5 !important;
  }}

  /* Selected / highlighted option */
  [aria-selected="true"][role="option"],
  [data-highlighted][role="option"] {{
    background-color: #D6E8F7 !important;
    color: #0077C5 !important;
  }}
  [aria-selected="true"][role="option"] *,
  [data-highlighted][role="option"] * {{
    color: #0077C5 !important;
  }}

  /* "No results" message */
  [data-baseweb="menu"] li[aria-disabled="true"],
  [data-baseweb="menu"] li[aria-disabled="true"] * {{
    color: #6B7F96 !important;
    background-color: #FFFFFF !important;
  }}
</style>
""", unsafe_allow_html=True)

# ── Column name constants ─────────────────────────────────────────────────────
R_PIPELINE_ID  = "Pipelines > Pipeline ID"
R_L2           = "DS - Join Supervisory Organization > LVL 2 Org Name"
R_L3           = "DS - Join Supervisory Organization > LVL 3 Org Name"
R_L4           = "DS - Join Supervisory Organization > LVL 4 Org Name"
R_JOB_LEVEL    = "Pipelines > Job level"
R_JOB_PROFILE  = "Pipelines > Job profile name"
R_JOB_TITLE    = "Pipelines > Job posting title"
R_JOB_FAMILY   = "Pipelines > Job family"
R_REMAINING    = "Pipelines > Remaining openings"
R_OPENINGS     = "Pipelines > # Openings"
R_STEP         = "Pipeline current step > Step"
R_TARGET_FY    = "Pipelines > Target fiscal year"
R_TARGET_QTR   = "Pipelines > Target quarter"
R_PRIORITY     = "Pipelines > Target priority level"
R_PRIORITY_CAT = "Pipelines > Target priority category"
R_RECRUITER    = "Pipelines > Recruiter"
R_TAM          = "Pipelines > TAM"
R_LOCATION     = "Pipelines > Work location"
R_COUNTRY      = "Pipelines > Location country"
R_DAYS_OPEN    = "Days Open"
R_JOB_FAM_GRP  = "Pipelines > Job family group"
R_WORKER_TYPE  = "Pipelines > Worker Type Classification - Employee Type"

def _resolve_col(df: pd.DataFrame, wanted: str) -> str:
    """Return the actual column name that best matches `wanted`.
    Tries exact match first, then case-insensitive exact.
    Falls back to `wanted` so downstream code surfaces a clean KeyError."""
    cols = list(df.columns)
    if wanted in cols:
        return wanted
    low = wanted.lower()
    for c in cols:
        if c.lower() == low:
            return c
    return wanted

def _find_org_col(df: pd.DataFrame, level: int) -> str:
    """Find the L2/L3/L4 org column regardless of how Workday exported it.
    Matches patterns like 'Level 02', 'LVL 2', 'L2', 'level 2' etc."""
    import re
    pattern = re.compile(
        rf"(level\s*0?{level}|lvl\s*0?{level})\b", re.IGNORECASE
    )
    for c in df.columns:
        if pattern.search(c):
            return c
    return None

H_PIPELINE_ID  = "Avature Pipeline ID"
H_L2           = "Org Level 2"
H_L3           = "Org Level 3"
H_L4           = "Org Level 4"
H_JOB_LEVEL    = "Job Level"
H_JOB_TITLE    = "Job Title"
H_JOB_FAMILY   = "Job Family"
H_JOB_FAM_GRP  = "Job Family Group"
H_START_DATE   = "Job Start Date"
H_HIRE_TYPE    = "Hire Type"
H_EMP_TYPE     = "Employee Type"
H_TARGET_FY    = "Target Fiscal Year"
H_TARGET_QTR   = "Target Quarter"
H_PRIORITY     = "Target Priority Level"
H_PRIORITY_CAT = "Target Priority Category"
H_RECRUITER    = "Recruiter Name"
H_TAM          = "TAM Name"
H_HM           = "Hiring Manager"
H_SITE         = "Business Site"
H_COUNTRY      = "Country"
H_CAREER_TRACK = "Career Track"
H_COMMUNITY    = "Community"
H_TECH_COMM    = "Tech Community"

# Actual Hires YTD column constants
A_NAME         = "Employee Name"
A_HIRE_DATE    = "Hire Date"
A_HIRE_TYPE    = "Hire Type"
A_BAND         = "Band"
A_MANAGER      = "Manager"
A_TECH_COMM    = "Tech Community"
A_COMMUNITY    = "Community"
A_JOB_CODE     = "Job Code"
A_JOB_LEVEL    = "Job Level"
A_JOB_FAMILY   = "Job Family"
A_JOB_FAM_GRP  = "Job Family Group"
A_JOB_TITLE    = "Job Title"
A_CAREER_TRACK = "Career Track"
A_MGT_LEVEL    = "Management Level"
A_L2           = "Org Level 2"
A_L3           = "Org Level 3"
A_L4           = "Org Level 4"
A_L5           = "Org Level 5"
A_SITE         = "Business Site"
A_STATE        = "State"
A_COUNTRY      = "Country"
A_AS_OF_DATE   = "As Of Date"
A_FISCAL_MONTH = "Fiscal Month"
A_FISCAL_YEAR  = "Fiscal Year"

# ── Data loading ──────────────────────────────────────────────────────────────
@st.cache_data(show_spinner=False)
def load_open_reqs(raw: bytes) -> pd.DataFrame:
    df = pd.read_excel(raw)
    df.columns = [c.strip() for c in df.columns]
    # Detect L2/L3/L4 org columns regardless of Workday export format
    # (handles "Level 02 Organization Name", "LVL 2 Org Name", etc.)
    rename = {}
    for level, const in [(2, R_L2), (3, R_L3), (4, R_L4)]:
        found = _find_org_col(df, level)
        if found and found != const:
            rename[found] = const
    if rename:
        df = df.rename(columns=rename)
    return df

@st.cache_data(show_spinner=False)
def load_expected_hires(raw: bytes) -> pd.DataFrame:
    df = pd.read_excel(raw)
    df.columns = [c.strip() for c in df.columns]
    df[H_START_DATE] = pd.to_datetime(df[H_START_DATE], errors="coerce")
    df["_month_sort"] = df[H_START_DATE].dt.to_period("M")
    df["Month"]       = df["_month_sort"].astype(str)
    df["Month Label"] = df[H_START_DATE].dt.strftime("%b %Y")
    return df

@st.cache_data(show_spinner=False)
def load_actual_hires(raw: bytes) -> pd.DataFrame:
    df = pd.read_excel(raw)
    df.columns = [c.strip() for c in df.columns]
    df[A_HIRE_DATE]   = pd.to_datetime(df[A_HIRE_DATE],   errors="coerce")
    df[A_AS_OF_DATE]  = pd.to_datetime(df[A_AS_OF_DATE],  errors="coerce")
    df["_month_sort"] = df[A_HIRE_DATE].dt.to_period("M")
    df["Month"]       = df["_month_sort"].astype(str)
    df["Month Label"] = df[A_HIRE_DATE].dt.strftime("%b %Y")
    return df

# ── Cross-filter state helpers ────────────────────────────────────────────────
def _xf_key(tab: str) -> str:
    return f"__xf_{tab}"

def get_xfilter(tab: str) -> list:
    """Return list of (l3, job_level) tuples currently selected."""
    return st.session_state.get(_xf_key(tab), [])

def toggle_xfilter(tab: str, l3: str, jl: str):
    """Add or remove an (l3, jl) selection; clear if already selected."""
    key  = _xf_key(tab)
    curr = st.session_state.get(key, [])
    pair = (l3, jl)
    if pair in curr:
        curr = [x for x in curr if x != pair]
    else:
        curr = curr + [pair]
    st.session_state[key] = curr

def apply_xfilter(df, tab: str, l3_col: str, jl_col: str) -> pd.DataFrame:
    """Filter df by selected (l3, jl) pairs. No selection = no filter."""
    sel = get_xfilter(tab)
    if not sel:
        return df
    mask = pd.Series(False, index=df.index)
    for l3, jl in sel:
        mask |= (df[l3_col] == l3) & (df[jl_col] == jl)
    return df[mask]

# ── Helpers ───────────────────────────────────────────────────────────────────
def ms_filter(df, col, sel):
    return df[df[col].isin(sel)] if sel else df

def clean_options(series: pd.Series) -> list:
    return sorted(series.replace("-", pd.NA).dropna().unique().tolist())

def pivot_with_totals(df, index_col, col_col, val_col, aggfunc="sum") -> pd.DataFrame:
    piv = df.pivot_table(index=index_col, columns=col_col,
                         values=val_col, aggfunc=aggfunc, fill_value=0)
    piv = piv.reindex(sorted(piv.columns), axis=1)
    piv["Total"] = piv.sum(axis=1)
    piv = piv.sort_values("Total", ascending=False)
    grand = piv.sum()
    grand.name = "Grand Total"
    return pd.concat([piv, grand.to_frame().T])

def simple_summary(df, group_col, val_col="count") -> pd.DataFrame:
    if val_col == "count":
        s = df.groupby(group_col).size().reset_index(name="Count")
    else:
        s = df.groupby(group_col)[val_col].sum().reset_index()
    total_row = pd.DataFrame({group_col: ["Total"], s.columns[-1]: [s.iloc[:, -1].sum()]})
    return pd.concat([s.sort_values(s.columns[-1], ascending=False), total_row], ignore_index=True)

_ALL = "✅ Select All"

def ms_with_all(label: str, options: list, key: str, default=None) -> list:
    """Multiselect with Select All support using on_change callback.

    The on_change fires BEFORE the next render, so when Select All is
    detected we swap the widget's own key value to the full list —
    no rerun needed, no key conflicts.

    Workflow:
      1. Open dropdown → pick '✅ Select All' → all options appear selected
      2. Click any blue pill tag to remove it
      3. Empty selection = no filter = show everything
    """
    options = list(options)
    choices = [_ALL] + options

    def _on_change():
        val = st.session_state[key]
        if _ALL in val:
            # Replace with all real options (no Select All pill in result)
            st.session_state[key] = options

    # Initialise default only on first render
    if key not in st.session_state:
        st.session_state[key] = list(default or [])

    st.multiselect(label, choices, key=key, on_change=_on_change)
    return st.session_state[key]

def chart_base():
    return dict(
        plot_bgcolor=WHITE,
        paper_bgcolor=WHITE,
        font=dict(family="Arial", color=DARK, size=12),
        margin=dict(t=40, b=20, l=10, r=10),
        xaxis=dict(gridcolor="#E8ECF0", linecolor=BORDER),
        yaxis=dict(gridcolor="#E8ECF0", linecolor=BORDER),
    )

# ── Sidebar ───────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown(f"<h3 style='color:{BLUE};margin-top:0'>Data Sources</h3>", unsafe_allow_html=True)
    st.caption("Drop files into the dashboard folder, then click Refresh to update.")

    # ── File status indicators ────────────────────────────────────────────────
    for label, path in [("open_reqs.xlsx", STD_REQS), ("expected_hires.xlsx", STD_HIRES), ("actual_hires.xlsx", STD_ACTUAL)]:
        if os.path.exists(path):
            mtime = datetime.fromtimestamp(os.path.getmtime(path)).strftime("%d %b %Y, %H:%M")
            st.markdown(f"✅ **{label}**  \n<small style='color:#888'>{mtime}</small>", unsafe_allow_html=True)
        else:
            st.markdown(f"⬜ **{label}**  \n<small style='color:#cc0000'>not found</small>", unsafe_allow_html=True)

    st.markdown("---")
    st.caption("After dropping new files into the dashboard folder, click Refresh to reload all data.")
    if st.button("🔄 Refresh Dashboard", use_container_width=True, type="primary", key="sidebar_refresh"):
        st.cache_data.clear()
        for key in ["reqs_bytes", "hires_bytes", "actual_hires_bytes"]:
            st.session_state.pop(key, None)
        st.rerun()

    st.markdown("---")
    st.markdown(f"<small style='color:#888'>Intuit Talent Acquisition · FY26</small>", unsafe_allow_html=True)

# Load data — always re-read from disk if the file is newer than what's in session
def _file_mtime_ts(path):
    """Return file mtime as a float, or 0 if file doesn't exist."""
    return os.path.getmtime(path) if os.path.exists(path) else 0

for _sess_key, _path, _cache_key in [
    ("reqs_bytes",         STD_REQS,   "reqs"),
    ("hires_bytes",        STD_HIRES,  "hires"),
    ("actual_hires_bytes", STD_ACTUAL, "actual_hires"),
]:
    _file_ts   = _file_mtime_ts(_path)
    _loaded_ts = st.session_state.get(f"{_sess_key}_mtime", 0)

    if _sess_key not in st.session_state or _file_ts > _loaded_ts:
        # File is new or has been updated — reload it
        data = load_from_folder(_path, _cache_key) or load_from_cache(_cache_key)
        if data:
            st.session_state[_sess_key]              = data
            st.session_state[f"{_sess_key}_mtime"]   = _file_ts
            st.cache_data.clear()

reqs_bytes         = st.session_state.get("reqs_bytes")
hires_bytes        = st.session_state.get("hires_bytes")
actual_hires_bytes = st.session_state.get("actual_hires_bytes")

# ── Header ────────────────────────────────────────────────────────────────────
hdr_col, refresh_col = st.columns([5, 1])
with hdr_col:
    st.markdown("<h1 style='margin-bottom:0'>Intuit Hiring Dashboard</h1>", unsafe_allow_html=True)
with refresh_col:
    st.markdown("<div style='padding-top:18px'>", unsafe_allow_html=True)
    if st.button("🔄 Refresh Data", use_container_width=True, type="primary", key="header_refresh"):
        st.cache_data.clear()
        for key in ["reqs_bytes", "hires_bytes", "actual_hires_bytes"]:
            st.session_state.pop(key, None)
        st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)

# Timestamp banner — uses folder file mtime if available, else falls back to cache meta
def file_mtime(path):
    return datetime.fromtimestamp(os.path.getmtime(path)).strftime("%d %b %Y, %H:%M") if os.path.exists(path) else None

meta = load_meta()

def best_ts(path, cache_key):
    return file_mtime(path) or meta.get(cache_key, "not loaded")

ts_reqs   = best_ts(STD_REQS,   "reqs")
ts_hires  = best_ts(STD_HIRES,  "hires")
ts_actual = best_ts(STD_ACTUAL, "actual_hires")

st.markdown(
    f"<div class='ts-banner'>"
    f"📁 <strong>Data last loaded</strong> &nbsp;|&nbsp; "
    f"Open Reqs: <strong>{ts_reqs}</strong> &nbsp;·&nbsp; "
    f"Expected Hires: <strong>{ts_hires}</strong> &nbsp;·&nbsp; "
    f"Actual Hires YTD: <strong>{ts_actual}</strong>"
    f"</div>",
    unsafe_allow_html=True,
)

if not reqs_bytes and not hires_bytes and not actual_hires_bytes:
    st.info("No data loaded yet. Drop **open_reqs.xlsx**, **expected_hires.xlsx**, and **actual_hires.xlsx** into the dashboard folder — or upload via the sidebar buttons.")
    st.stop()

df_reqs         = load_open_reqs(reqs_bytes)             if reqs_bytes         else None
df_hires        = load_expected_hires(hires_bytes)       if hires_bytes        else None
df_actual_hires = load_actual_hires(actual_hires_bytes)  if actual_hires_bytes else None

# Pipeline IDs that have expected hires
pipeline_ids_with_offers: set = set()
if df_hires is not None:
    pipeline_ids_with_offers = set(df_hires[H_PIPELINE_ID].dropna().astype(int).unique())

if df_reqs is not None:
    df_reqs["Has Expected Offer"] = (
        df_reqs[R_PIPELINE_ID].isin(pipeline_ids_with_offers)
        .map({True: "Yes", False: "No"})
    )

# ══════════════════════════════════════════════════════════════════════════════
tab1, tab2, tab3, tab4 = st.tabs(["  Open Hiring Requisitions  ", "  Expected Hires  ", "  Actual Hires YTD  ", "  Reports  "])

# ─────────────────────────────────────────────────────────────────────────────
# TAB 1
# ─────────────────────────────────────────────────────────────────────────────
with tab1:
    if df_reqs is None:
        st.warning("Upload the **Open Hiring Requisitions** file in the sidebar.")
    else:
        # ── Filters ──────────────────────────────────────────────────────────
        st.markdown("<div class='filter-card'>", unsafe_allow_html=True)
        st.markdown("**Filters**")

        # Row 1 — Employee Type (always first, default = Employee) + Org hierarchy
        fr1a, fr1b, fr1c, fr1d = st.columns(4)
        with fr1a:
            wt_opts = clean_options(df_reqs[R_WORKER_TYPE])
            wt_default = ["Employee"] if "Employee" in wt_opts else []
            wt_filter = ms_with_all("Employee Type", wt_opts, key="r_wt", default=wt_default)
        with fr1b:
            l2_filter = ms_with_all("Level 2 Org", clean_options(df_reqs[R_L2]), key="r_l2")
        with fr1c:
            l3_filter = ms_with_all("Level 3 Org", clean_options(df_reqs[R_L3]), key="r_l3")
        with fr1d:
            l4_filter = ms_with_all("Level 4 Org", clean_options(df_reqs[R_L4]), key="r_l4")

        # Row 2 — Job dimensions
        fr2a, fr2b, fr2c, fr2d = st.columns(4)
        with fr2a:
            jfg_filter = ms_with_all("Job Family Group", clean_options(df_reqs[R_JOB_FAM_GRP]), key="r_jfg")
        with fr2b:
            jf_filter  = ms_with_all("Job Family",       clean_options(df_reqs[R_JOB_FAMILY]),  key="r_jf")
        with fr2c:
            jl_filter  = ms_with_all("Job Level",        clean_options(df_reqs[R_JOB_LEVEL]),   key="r_jl")
        with fr2d:
            jp_filter  = ms_with_all("Job Profile",      clean_options(df_reqs[R_JOB_PROFILE]), key="r_jp")

        # Row 3 — People & Priority
        fr3a, fr3b, fr3c, fr3d = st.columns(4)
        with fr3a:
            tam_filter    = ms_with_all("TAM",               clean_options(df_reqs[R_TAM]),               key="r_tam")
        with fr3b:
            rec_filter    = ms_with_all("Recruiter",         clean_options(df_reqs[R_RECRUITER]),         key="r_rec")
        with fr3c:
            pri_filter    = ms_with_all("Priority Level",    clean_options(df_reqs[R_PRIORITY]),          key="r_pri")
        with fr3d:
            pricat_filter = ms_with_all("Priority Category", clean_options(df_reqs[R_PRIORITY_CAT]),      key="r_pricat")

        # Row 4 — Step, FY, Offer toggle, Pipeline IDs
        fr4a, fr4b, fr4c, fr4d = st.columns(4)
        with fr4a:
            step_filter = ms_with_all("Pipeline Step", clean_options(df_reqs[R_STEP]),       key="r_step")
        with fr4b:
            fy_filter   = ms_with_all("Target FY",    clean_options(df_reqs[R_TARGET_FY]),   key="r_fy")
        with fr4c:
            offer_toggle = st.selectbox(
                "Has Expected Offer?", ["All", "Yes — has offer", "No — no offer"], key="r_offer"
            )
        with fr4d:
            if df_hires is not None and offer_toggle != "No — no offer":
                pid_filter = ms_with_all(
                    "Specific Pipeline IDs",
                    sorted(pipeline_ids_with_offers), key="r_pid",
                )
            else:
                pid_filter = []

        st.markdown("</div>", unsafe_allow_html=True)

        # Apply filters
        f = df_reqs.copy()
        f = ms_filter(f, R_WORKER_TYPE,  wt_filter)
        f = ms_filter(f, R_L2,           l2_filter)
        f = ms_filter(f, R_L3,           l3_filter)
        f = ms_filter(f, R_L4,           l4_filter)
        f = ms_filter(f, R_JOB_FAM_GRP,  jfg_filter)
        f = ms_filter(f, R_JOB_FAMILY,   jf_filter)
        f = ms_filter(f, R_JOB_LEVEL,    jl_filter)
        f = ms_filter(f, R_JOB_PROFILE,  jp_filter)
        f = ms_filter(f, R_TAM,          tam_filter)
        f = ms_filter(f, R_RECRUITER,    rec_filter)
        f = ms_filter(f, R_PRIORITY,     pri_filter)
        f = ms_filter(f, R_PRIORITY_CAT, pricat_filter)
        f = ms_filter(f, R_STEP,         step_filter)
        f = ms_filter(f, R_TARGET_FY,    fy_filter)
        if offer_toggle == "Yes — has offer":
            f = f[f["Has Expected Offer"] == "Yes"]
        elif offer_toggle == "No — no offer":
            f = f[f["Has Expected Offer"] == "No"]
        if pid_filter:
            f = f[f[R_PIPELINE_ID].isin(pid_filter)]

        # ── Metrics ──────────────────────────────────────────────────────────
        m1, m2, m3, m4, m5 = st.columns(5)
        m1.metric("Total Requisitions",     f"{len(f):,}")
        m2.metric("Remaining Openings",     f"{int(f[R_REMAINING].sum()):,}")
        m3.metric("With Expected Offer",    f"{(f['Has Expected Offer']=='Yes').sum():,}")
        m4.metric("Without Expected Offer", f"{(f['Has Expected Offer']=='No').sum():,}")
        m5.metric("Avg Days Open",
                  f"{f[R_DAYS_OPEN].mean():.0f}" if f[R_DAYS_OPEN].notna().any() else "—")

        st.markdown("---")

        # ── Table 1: L2 Org — 3-column split ─────────────────────────────────
        st.markdown("<p class='section-label'>Remaining Openings by Level 2 Org</p>", unsafe_allow_html=True)
        l2_grp = (
            f.groupby([R_L2, "Has Expected Offer"])[R_REMAINING].sum()
            .unstack(fill_value=0)
        )
        for col in ["Yes", "No"]:
            if col not in l2_grp.columns:
                l2_grp[col] = 0
        l2_grp = l2_grp.rename(columns={"Yes": "With Offer", "No": "Without Offer"})
        l2_grp["Total Remaining"] = l2_grp["With Offer"] + l2_grp["Without Offer"]
        l2_grp = l2_grp[["With Offer", "Without Offer", "Total Remaining"]].sort_values("Total Remaining", ascending=False)
        grand_l2 = l2_grp.sum(); grand_l2.name = "Total"
        l2_grp = pd.concat([l2_grp, grand_l2.to_frame().T])
        l2_grp.index.name = "Level 2 Org"
        st.dataframe(l2_grp, use_container_width=True)

        st.markdown("---")

        # ── Table 2: L3 Org × Job Level — clickable cross-filter heatmap ───────
        st.markdown("<p class='section-label'>Remaining Openings by L3 Org × Job Level</p>", unsafe_allow_html=True)

        xf_sel = get_xfilter("tab1")
        if xf_sel:
            sel_l3s = sorted({l3 for l3, _ in xf_sel})
            sel_jls = sorted({jl for _, jl in xf_sel})
            st.markdown(
                f"<div style='background:#FFF3CD;border-left:4px solid #FF6900;"
                f"border-radius:4px;padding:8px 14px;margin-bottom:8px;font-size:0.85rem;color:#1C2B3A'>"
                f"🔍 <strong>Cross-filter active</strong> — "
                f"L3: <strong>{', '.join(sel_l3s)}</strong> &nbsp;·&nbsp; "
                f"Level: <strong>{', '.join(sel_jls)}</strong> &nbsp;·&nbsp; "
                f"<em>Click a cell again to deselect, or clear all below</em></div>",
                unsafe_allow_html=True,
            )
            if st.button("✖ Clear cross-filter", key="xf_clear_tab1"):
                st.session_state[_xf_key("tab1")] = []
                st.rerun()

        l3_view = st.radio(
            "Show openings:",
            ["Total", "With Offer", "Without Offer"],
            horizontal=True, key="l3_jl_toggle",
        )
        if l3_view == "With Offer":
            l3_src = f[f["Has Expected Offer"] == "Yes"]
        elif l3_view == "Without Offer":
            l3_src = f[f["Has Expected Offer"] == "No"]
        else:
            l3_src = f

        # Build pivot for heatmap (exclude Grand Total row for clicking)
        _piv_raw = l3_src.pivot_table(
            index=R_L3, columns=R_JOB_LEVEL, values=R_REMAINING,
            aggfunc="sum", fill_value=0
        )
        _piv_raw = _piv_raw.reindex(sorted(_piv_raw.columns), axis=1)
        _piv_raw["Total"] = _piv_raw.sum(axis=1)
        _piv_raw = _piv_raw.sort_values("Total", ascending=False)

        _jl_cols  = [c for c in _piv_raw.columns if c != "Total"]
        _l3_rows  = list(_piv_raw.index)
        _z        = _piv_raw[_jl_cols].values.tolist()
        _text     = [[str(int(v)) if v else "" for v in row] for row in _z]

        # Highlight selected cells
        _sel_set  = set(xf_sel)
        _colors   = [
            ["#0077C5" if (_l3_rows[r], _jl_cols[c]) in _sel_set else
             ("#E8F0F8" if _z[r][c] > 0 else "#F7F9FB")
             for c in range(len(_jl_cols))]
            for r in range(len(_l3_rows))
        ]
        _font_col = [
            ["#FFFFFF" if (_l3_rows[r], _jl_cols[c]) in _sel_set else "#1C2B3A"
             for c in range(len(_jl_cols))]
            for r in range(len(_l3_rows))
        ]

        fig_heat = go.Figure(go.Heatmap(
            z=_z, x=_jl_cols, y=_l3_rows,
            text=_text, texttemplate="%{text}",
            colorscale=[[0, "#F7F9FB"], [1, "#0077C5"]],
            showscale=False,
            hovertemplate="<b>%{y}</b><br>%{x}: %{z}<extra></extra>",
        ))
        # Overlay coloured cells via shapes + annotations for selected
        for r, l3 in enumerate(_l3_rows):
            for c, jl in enumerate(_jl_cols):
                if (l3, jl) in _sel_set:
                    fig_heat.add_shape(
                        type="rect",
                        x0=c - 0.5, x1=c + 0.5,
                        y0=r - 0.5, y1=r + 0.5,
                        fillcolor="#FF6900", opacity=0.85,
                        line=dict(color="#FF6900"),
                    )

        fig_heat.update_layout(
            **chart_base(),
            height=max(280, len(_l3_rows) * 36 + 80),
            margin=dict(t=20, b=20, l=10, r=10),
            xaxis=dict(side="top", tickfont=dict(size=11, color="#1C2B3A")),
            yaxis=dict(tickfont=dict(size=11, color="#1C2B3A"), autorange="reversed"),
            clickmode="event",
        )
        hm_event = st.plotly_chart(
            fig_heat, use_container_width=True,
            key="xf_heatmap", on_select="rerun",
        )

        # Handle click — toggle selection
        if hm_event and hm_event.get("selection", {}).get("points"):
            for pt in hm_event["selection"]["points"]:
                clicked_jl = pt.get("x")
                clicked_l3 = pt.get("y")
                if clicked_l3 and clicked_jl and clicked_jl in _jl_cols:
                    toggle_xfilter("tab1", clicked_l3, clicked_jl)
                    # Sync L3 and Job Level dropdowns
                    new_l3s = sorted({l3 for l3, _ in get_xfilter("tab1")})
                    new_jls = sorted({jl for _, jl in get_xfilter("tab1")})
                    st.session_state["r_l3"]     = new_l3s
                    st.session_state["__r_l3_val"] = new_l3s
                    st.session_state["r_jl"]     = new_jls
                    st.session_state["__r_jl_val"] = new_jls
                    st.rerun()

        # Also show numeric table with totals below heatmap
        grand_row = _piv_raw.sum()
        grand_row.name = "Grand Total"
        _piv_display = pd.concat([_piv_raw, grand_row.to_frame().T])
        _piv_display.index.name = "L3 Org"
        st.dataframe(_piv_display, use_container_width=True)

        st.markdown("---")

        # ── Cross-filtered data for all tables below ──────────────────────────
        # fx = f filtered by clicked L3 × Job Level intersections
        fx = apply_xfilter(f, "tab1", R_L3, R_JOB_LEVEL)

        # ── Table 3: Job Level × Has Expected Offer (replaces chart) ─────────
        st.markdown("<p class='section-label'>Remaining Openings by Job Level — Offer Status</p>", unsafe_allow_html=True)

        jl_offer_piv = pivot_with_totals(fx, R_JOB_LEVEL, "Has Expected Offer", R_REMAINING)
        jl_offer_piv.index.name = "Job Level"
        # Rename columns for clarity
        jl_offer_piv = jl_offer_piv.rename(columns={"No": "No Offer", "Yes": "Has Offer"})

        tc1, tc2 = st.columns([2, 3])
        with tc1:
            st.dataframe(jl_offer_piv, use_container_width=True)
        with tc2:
            chart_src = (
                fx.groupby([R_JOB_LEVEL, "Has Expected Offer"])[R_REMAINING]
                .sum().reset_index()
                .rename(columns={R_JOB_LEVEL: "Job Level", R_REMAINING: "Openings", "Has Expected Offer": "Offer"})
            )
            fig = px.bar(
                chart_src, x="Job Level", y="Openings", color="Offer",
                color_discrete_map={"Yes": GREEN, "No": ORANGE},
                barmode="stack",
                title="Job Level — Offer Status Split",
            )
            fig.update_layout(**chart_base(), height=320,
                              legend=dict(orientation="h", y=-0.2, title=""),
                              title_font=dict(size=13))
            st.plotly_chart(fig, use_container_width=True)

        st.markdown("---")

        # ── Table 4: Job Family Group × Job Level — Offer Status ─────────────
        st.markdown("<p class='section-label'>Remaining Openings by Job Family Group & Job Level — Offer Status</p>", unsafe_allow_html=True)

        jfg_view = st.radio(
            "Show openings:",
            ["Total", "Has Offer", "No Offer"],
            horizontal=True, key="jfg_jl_toggle",
        )
        if jfg_view == "Has Offer":
            jfg_src = fx[fx["Has Expected Offer"] == "Yes"]
        elif jfg_view == "No Offer":
            jfg_src = fx[fx["Has Expected Offer"] == "No"]
        else:
            jfg_src = fx

        if R_JOB_FAM_GRP in jfg_src.columns:
            jfg_jl_piv = pivot_with_totals(jfg_src, R_JOB_FAM_GRP, R_JOB_LEVEL, R_REMAINING)
            jfg_jl_piv.index.name = "Job Family Group"
            st.dataframe(jfg_jl_piv, use_container_width=True)
        else:
            st.info("Job Family Group column not found in this data export.")

        st.markdown("---")

        # ── Table 5: Priority breakdown ──────────────────────────────────────
        st.markdown("<p class='section-label'>Remaining Openings by Priority Level</p>", unsafe_allow_html=True)
        pri_tbl = (
            fx.dropna(subset=[R_PRIORITY])
            .groupby(R_PRIORITY)[R_REMAINING].sum()
            .reset_index().rename(columns={R_PRIORITY: "Priority", R_REMAINING: "Remaining Openings"})
            .sort_values("Priority")
        )
        pri_total = pd.DataFrame({"Priority": ["Total"], "Remaining Openings": [pri_tbl["Remaining Openings"].sum()]})
        st.dataframe(
            pd.concat([pri_tbl, pri_total], ignore_index=True),
            use_container_width=True, hide_index=True,
        )

        st.markdown("---")

        # ── Main detail table ─────────────────────────────────────────────────
        xf_label = f" · filtered to {len(fx):,} rows by cross-filter" if xf_sel else ""
        st.markdown("<p class='section-label'>All Open Requisitions</p>", unsafe_allow_html=True)
        st.caption(f"Sorted: L3 Org → L4 Org → Job Level  ·  {len(fx):,} rows{xf_label}")

        col_map = {
            R_PIPELINE_ID:  "Pipeline ID",
            R_WORKER_TYPE:  "Employee Type",
            R_L2:           "L2 Org",
            R_L3:           "L3 Org",
            R_L4:           "L4 Org",
            R_JOB_FAM_GRP:  "Job Family Group",
            R_JOB_FAMILY:   "Job Family",
            R_JOB_LEVEL:    "Job Level",
            R_JOB_PROFILE:  "Job Profile",
            R_JOB_TITLE:    "Job Title",
            R_REMAINING:    "Remaining",
            R_OPENINGS:     "# Openings",
            R_STEP:         "Current Step",
            R_TARGET_FY:    "Target FY",
            R_TARGET_QTR:   "Quarter",
            R_PRIORITY:     "Priority",
            R_PRIORITY_CAT: "Priority Category",
            R_RECRUITER:    "Recruiter",
            R_TAM:          "TAM",
            R_LOCATION:     "Location",
            R_DAYS_OPEN:    "Days Open",
            "Has Expected Offer": "Has Offer?",
        }
        tbl = (
            fx[[c for c in col_map if c in fx.columns]]
            .rename(columns=col_map)
            .sort_values(["L3 Org", "L4 Org", "Job Level"])
            .reset_index(drop=True)
        )

        def row_style(row):
            if row.get("Has Offer?") == "Yes":
                return ["background-color:#EAF5EA; color:#1A1A2E"] * len(row)
            return [""] * len(row)

        st.dataframe(
            tbl.style.apply(row_style, axis=1),
            use_container_width=True, height=500,
        )
        st.download_button(
            "Download Filtered Data (CSV)",
            tbl.to_csv(index=False).encode("utf-8"),
            "open_reqs_filtered.csv", "text/csv",
        )

# ─────────────────────────────────────────────────────────────────────────────
# TAB 2
# ─────────────────────────────────────────────────────────────────────────────
with tab2:
    if df_hires is None:
        st.warning("Upload the **Expected Hires** file in the sidebar.")
    else:
        # ── Filters ──────────────────────────────────────────────────────────
        # ── Start-date toggle ─────────────────────────────────────────────────
        today = datetime.now().date()
        st.markdown("<div class='filter-card'>", unsafe_allow_html=True)
        h_start_toggle = st.radio(
            "Start Date Status",
            ["All", "Started (start date before today)", "Starting on/after today"],
            horizontal=True,
            key="h_start_toggle",
            help=f"Today's date: {today.strftime('%d %b %Y')}",
        )

        st.markdown("**Filters**")

        # Row 1 — Employee Type (first, default Employee) + Org hierarchy
        hf1a, hf1b, hf1c, hf1d = st.columns(4)
        with hf1a:
            het_opts    = clean_options(df_hires[H_EMP_TYPE]) if H_EMP_TYPE in df_hires.columns else []
            het_default = ["Employee"] if "Employee" in het_opts else []
            het = ms_with_all("Employee Type", het_opts, key="h_et", default=het_default)
        with hf1b:
            hl2  = ms_with_all("Org Level 2", clean_options(df_hires[H_L2])  if H_L2  in df_hires.columns else [], key="h_l2")
        with hf1c:
            hl3  = ms_with_all("Org Level 3", clean_options(df_hires[H_L3]),  key="h_l3")
        with hf1d:
            hl4  = ms_with_all("Org Level 4", clean_options(df_hires[H_L4])  if H_L4  in df_hires.columns else [], key="h_l4")

        # Row 2 — Job dimensions
        hf2a, hf2b, hf2c, hf2d = st.columns(4)
        with hf2a:
            hjfg = ms_with_all("Job Family Group", clean_options(df_hires[H_JOB_FAM_GRP]) if H_JOB_FAM_GRP in df_hires.columns else [], key="h_jfg")
        with hf2b:
            hjf  = ms_with_all("Job Family",       clean_options(df_hires[H_JOB_FAMILY]),  key="h_jf")
        with hf2c:
            hjl  = ms_with_all("Job Level",        clean_options(df_hires[H_JOB_LEVEL]),   key="h_jl")
        with hf2d:
            hht  = ms_with_all("Hire Type",        clean_options(df_hires[H_HIRE_TYPE]),   key="h_ht")

        # Row 3 — People & Priority
        hf3a, hf3b, hf3c, hf3d = st.columns(4)
        with hf3a:
            htam = ms_with_all("TAM",               clean_options(df_hires[H_TAM]),                                                              key="h_tam")
        with hf3b:
            hrec = ms_with_all("Recruiter",         clean_options(df_hires[H_RECRUITER]),                                                        key="h_rec")
        with hf3c:
            hpp  = ms_with_all("Priority Level",    clean_options(df_hires[H_PRIORITY]),                                                         key="h_pri")
        with hf3d:
            hppc = ms_with_all("Priority Category", clean_options(df_hires[H_PRIORITY_CAT]) if H_PRIORITY_CAT in df_hires.columns else [],       key="h_pricat")

        # Row 4 — FY, Quarter & Start Month
        hf4a, hf4b, hf4c, _ = st.columns(4)
        with hf4a:
            hfy  = ms_with_all("Target FY", clean_options(df_hires[H_TARGET_FY]),  key="h_fy")
        with hf4b:
            hqtr = ms_with_all("Quarter",   clean_options(df_hires[H_TARGET_QTR]), key="h_qtr")
        with hf4c:
            _month_opts = (
                df_hires.dropna(subset=["_month_sort"])
                .drop_duplicates(subset=["Month Label"])
                .sort_values("_month_sort")["Month Label"]
                .tolist()
            )
            hmonth = ms_with_all("Start Month", _month_opts, key="h_month")

        # Row 5 — Career Track / Community / Tech Community
        hf5a, hf5b, hf5c, _ = st.columns(4)
        with hf5a:
            hct  = ms_with_all("Career Track",   clean_options(df_hires[H_CAREER_TRACK]) if H_CAREER_TRACK in df_hires.columns else [], key="h_ct")
        with hf5b:
            hcom = ms_with_all("Community",      clean_options(df_hires[H_COMMUNITY])    if H_COMMUNITY    in df_hires.columns else [], key="h_com")
        with hf5c:
            htc  = ms_with_all("Tech Community", clean_options(df_hires[H_TECH_COMM])    if H_TECH_COMM    in df_hires.columns else [], key="h_tc")

        st.markdown("</div>", unsafe_allow_html=True)

        fh = df_hires.copy()
        if het  and H_EMP_TYPE     in fh.columns: fh = ms_filter(fh, H_EMP_TYPE,     het)
        if hl2  and H_L2           in fh.columns: fh = ms_filter(fh, H_L2,           hl2)
        fh = ms_filter(fh, H_L3,         hl3)
        if hl4  and H_L4           in fh.columns: fh = ms_filter(fh, H_L4,           hl4)
        if hjfg and H_JOB_FAM_GRP  in fh.columns: fh = ms_filter(fh, H_JOB_FAM_GRP,  hjfg)
        fh = ms_filter(fh, H_JOB_FAMILY,  hjf)
        fh = ms_filter(fh, H_JOB_LEVEL,   hjl)
        fh = ms_filter(fh, H_HIRE_TYPE,   hht)
        fh = ms_filter(fh, H_TAM,         htam)
        fh = ms_filter(fh, H_RECRUITER,   hrec)
        fh = ms_filter(fh, H_PRIORITY,    hpp)
        if hppc and H_PRIORITY_CAT in fh.columns: fh = ms_filter(fh, H_PRIORITY_CAT, hppc)
        fh = ms_filter(fh, H_TARGET_FY,   hfy)
        fh = ms_filter(fh, H_TARGET_QTR,  hqtr)
        if hct  and H_CAREER_TRACK in fh.columns: fh = ms_filter(fh, H_CAREER_TRACK, hct)
        if hcom and H_COMMUNITY    in fh.columns: fh = ms_filter(fh, H_COMMUNITY,    hcom)
        if htc  and H_TECH_COMM    in fh.columns: fh = ms_filter(fh, H_TECH_COMM,    htc)
        if hmonth: fh = fh[fh["Month Label"].isin(hmonth)]

        if h_start_toggle == "Started (start date before today)":
            fh = fh[fh[H_START_DATE].dt.date < today]
        elif h_start_toggle == "Starting on/after today":
            fh = fh[fh[H_START_DATE].dt.date >= today]

        # ── Metrics ──────────────────────────────────────────────────────────
        hm1, hm2, hm3, hm4 = st.columns(4)
        hm1.metric("Total Expected Hires", f"{len(fh):,}")
        hm2.metric("Unique Pipelines",     f"{fh[H_PIPELINE_ID].nunique():,}")
        hm3.metric("Months Covered",       f"{fh['Month'].nunique():,}")
        hm4.metric("L3 Orgs",              f"{fh[H_L3].nunique():,}")

        st.markdown("---")

        # ── Monthly trend chart ───────────────────────────────────────────────
        st.markdown("<p class='section-label'>Expected Hires by Month</p>", unsafe_allow_html=True)
        monthly = (
            fh.groupby(["Month", "Month Label"]).size()
            .reset_index(name="Hires").sort_values("Month")
        )
        fig_trend = px.bar(monthly, x="Month Label", y="Hires",
                           color_discrete_sequence=[BLUE])
        fig_trend.update_traces(marker_line_color=WHITE, marker_line_width=1)
        fig_trend.update_layout(**chart_base(), height=280,
                                xaxis_title="", yaxis_title="# Expected Hires",
                                bargap=0.25)
        st.plotly_chart(fig_trend, use_container_width=True)

        st.markdown("---")

        # ── Table 1: Month × L3 Org ───────────────────────────────────────────
        st.markdown("<p class='section-label'>Expected Hires by Month × L3 Org</p>", unsafe_allow_html=True)
        piv_l3 = fh.pivot_table(index=H_L3, columns="Month", aggfunc="size", fill_value=0)
        piv_l3 = piv_l3.reindex(sorted(piv_l3.columns), axis=1)
        piv_l3["Total"] = piv_l3.sum(axis=1)
        piv_l3 = piv_l3.sort_values("Total", ascending=False)
        grand_l3 = piv_l3.sum(); grand_l3.name = "Grand Total"
        piv_l3 = pd.concat([piv_l3, grand_l3.to_frame().T])
        piv_l3.index.name = "L3 Org"
        st.dataframe(piv_l3, use_container_width=True)

        st.markdown("---")

        # ── Table 2: Month × Job Level ────────────────────────────────────────
        st.markdown("<p class='section-label'>Expected Hires by Month × Job Level</p>", unsafe_allow_html=True)
        piv_jl = fh.pivot_table(index=H_JOB_LEVEL, columns="Month", aggfunc="size", fill_value=0)
        piv_jl = piv_jl.reindex(sorted(piv_jl.columns), axis=1)
        piv_jl["Total"] = piv_jl.sum(axis=1)
        piv_jl = piv_jl.sort_values("Total", ascending=False)
        grand_jl = piv_jl.sum(); grand_jl.name = "Grand Total"
        piv_jl = pd.concat([piv_jl, grand_jl.to_frame().T])
        piv_jl.index.name = "Job Level"
        st.dataframe(piv_jl, use_container_width=True)

        st.markdown("---")

        # ── Table 3: L3 Org × Job Level cross-tab ────────────────────────────
        st.markdown("<p class='section-label'>Expected Hires by L3 Org × Job Level</p>", unsafe_allow_html=True)
        piv_l3jl = fh.pivot_table(index=H_L3, columns=H_JOB_LEVEL, aggfunc="size", fill_value=0)
        piv_l3jl["Total"] = piv_l3jl.sum(axis=1)
        piv_l3jl = piv_l3jl.sort_values("Total", ascending=False)
        grand_l3jl = piv_l3jl.sum(); grand_l3jl.name = "Grand Total"
        piv_l3jl = pd.concat([piv_l3jl, grand_l3jl.to_frame().T])
        piv_l3jl.index.name = "L3 Org"
        st.dataframe(piv_l3jl, use_container_width=True)

        st.markdown("---")

        # ── Table 4: Hire Type summary ────────────────────────────────────────
        st.markdown("<p class='section-label'>Expected Hires by Hire Type & Priority</p>", unsafe_allow_html=True)
        tc1, tc2 = st.columns(2)
        with tc1:
            ht_tbl = simple_summary(fh.replace("-", pd.NA).dropna(subset=[H_HIRE_TYPE]), H_HIRE_TYPE)
            ht_tbl.columns = ["Hire Type", "Count"]
            st.dataframe(ht_tbl, use_container_width=True, hide_index=True)
        with tc2:
            pri_h_tbl = simple_summary(fh.replace("-", pd.NA).dropna(subset=[H_PRIORITY]), H_PRIORITY)
            pri_h_tbl.columns = ["Priority", "Count"]
            st.dataframe(pri_h_tbl, use_container_width=True, hide_index=True)

        st.markdown("---")

        # ── Detail table ──────────────────────────────────────────────────────
        st.markdown("<p class='section-label'>Expected Hires Detail</p>", unsafe_allow_html=True)
        st.caption(f"Sorted: L3 Org → Job Level → Month  ·  {len(fh):,} rows")

        hire_col_map = {
            H_PIPELINE_ID:  "Pipeline ID",
            H_L3:           "L3 Org",
            H_L4:           "L4 Org",
            H_JOB_LEVEL:    "Job Level",
            H_JOB_TITLE:    "Job Title",
            H_JOB_FAMILY:   "Job Family",
            "Month Label":  "Start Month",
            H_HIRE_TYPE:    "Hire Type",
            H_TARGET_FY:    "Target FY",
            H_TARGET_QTR:   "Quarter",
            H_PRIORITY:     "Priority",
            H_PRIORITY_CAT: "Priority Category",
            H_RECRUITER:    "Recruiter",
            H_TAM:          "TAM",
            H_HM:           "Hiring Manager",
            H_SITE:         "Site",
            H_COUNTRY:      "Country",
        }
        hcols = [c for c in hire_col_map if c in fh.columns]
        hire_tbl = (
            fh[hcols].rename(columns=hire_col_map)
            .sort_values(["L3 Org", "Job Level", "Start Month"])
            .reset_index(drop=True)
        )
        st.dataframe(hire_tbl, use_container_width=True, height=480)
        st.download_button(
            "Download Expected Hires (CSV)",
            hire_tbl.to_csv(index=False).encode("utf-8"),
            "expected_hires_filtered.csv", "text/csv",
        )

# ─────────────────────────────────────────────────────────────────────────────
# TAB 3 — Actual Hires YTD
# ─────────────────────────────────────────────────────────────────────────────
with tab3:
    if df_actual_hires is None:
        st.warning("Upload the **Actual Hires YTD** file in the sidebar.")
    else:
        # Derive the upload/as-of date from the file (use max As Of Date)
        as_of_date = df_actual_hires[A_AS_OF_DATE].max()

        st.markdown("<div class='filter-card'>", unsafe_allow_html=True)
        st.markdown("**Filters**")

        # Row 1 — Org hierarchy
        af1a, af1b, af1c, af1d = st.columns(4)
        with af1a:
            al2   = ms_with_all("Org Level 2", clean_options(df_actual_hires[A_L2]),   key="a_l2")
        with af1b:
            al3   = ms_with_all("Org Level 3", clean_options(df_actual_hires[A_L3]),   key="a_l3")
        with af1c:
            al4   = ms_with_all("Org Level 4", clean_options(df_actual_hires[A_L4]),   key="a_l4")
        with af1d:
            aband = ms_with_all("Band",         clean_options(df_actual_hires[A_BAND]), key="a_band")

        # Row 2 — Job dimensions
        af2a, af2b, af2c, af2d = st.columns(4)
        with af2a:
            ajfg = ms_with_all("Job Family Group", clean_options(df_actual_hires[A_JOB_FAM_GRP]), key="a_jfg")
        with af2b:
            ajf  = ms_with_all("Job Family",       clean_options(df_actual_hires[A_JOB_FAMILY]),  key="a_jf")
        with af2c:
            ajl  = ms_with_all("Job Level",        clean_options(df_actual_hires[A_JOB_LEVEL]),   key="a_jl")
        with af2d:
            aht  = ms_with_all("Hire Type",        clean_options(df_actual_hires[A_HIRE_TYPE]),   key="a_ht")

        # Row 3 — People & Track
        af3a, af3b, af3c, af3d = st.columns(4)
        with af3a:
            act  = ms_with_all("Career Track",   clean_options(df_actual_hires[A_CAREER_TRACK]), key="a_ct")
        with af3b:
            acom = ms_with_all("Community",      clean_options(df_actual_hires[A_COMMUNITY]),    key="a_com")
        with af3c:
            atc  = ms_with_all("Tech Community", clean_options(df_actual_hires[A_TECH_COMM]),    key="a_tc")
        with af3d:
            afy  = ms_with_all("Fiscal Year",    clean_options(df_actual_hires[A_FISCAL_YEAR]),  key="a_fy")

        st.markdown("</div>", unsafe_allow_html=True)

        fa = df_actual_hires.copy()

        # Apply filters
        fa = ms_filter(fa, A_L2,           al2)
        fa = ms_filter(fa, A_L3,           al3)
        fa = ms_filter(fa, A_L4,           al4)
        fa = ms_filter(fa, A_BAND,         aband)
        fa = ms_filter(fa, A_JOB_FAM_GRP,  ajfg)
        fa = ms_filter(fa, A_JOB_FAMILY,   ajf)
        fa = ms_filter(fa, A_JOB_LEVEL,    ajl)
        fa = ms_filter(fa, A_HIRE_TYPE,    aht)
        fa = ms_filter(fa, A_CAREER_TRACK, act)
        fa = ms_filter(fa, A_COMMUNITY,    acom)
        fa = ms_filter(fa, A_TECH_COMM,    atc)
        fa = ms_filter(fa, A_FISCAL_YEAR,  afy)

        # ── Metrics ───────────────────────────────────────────────────────────
        already_started = (fa[A_HIRE_DATE] < as_of_date).sum() if pd.notna(as_of_date) else len(fa)
        upcoming        = (fa[A_HIRE_DATE] >= as_of_date).sum() if pd.notna(as_of_date) else 0
        am1, am2, am3, am4 = st.columns(4)
        am1.metric("Total Hires YTD",   f"{len(fa):,}")
        am2.metric("Already Started",   f"{already_started:,}")
        am3.metric("Starting Soon",      f"{upcoming:,}")
        am4.metric("L3 Orgs",           f"{fa[A_L3].nunique():,}")

        st.markdown("---")

        # ── Monthly trend ─────────────────────────────────────────────────────
        st.markdown("<p class='section-label'>Actual Hires by Month</p>", unsafe_allow_html=True)
        a_monthly = (
            fa.groupby(["Month", "Month Label"]).size()
            .reset_index(name="Hires").sort_values("Month")
        )
        fig_a_trend = px.bar(a_monthly, x="Month Label", y="Hires",
                             color_discrete_sequence=[GREEN])
        fig_a_trend.update_traces(marker_line_color=WHITE, marker_line_width=1)
        fig_a_trend.update_layout(**chart_base(), height=280,
                                  xaxis_title="", yaxis_title="# Actual Hires",
                                  bargap=0.25)
        st.plotly_chart(fig_a_trend, use_container_width=True)

        st.markdown("---")

        # ── Table 1: Month × L3 Org ───────────────────────────────────────────
        st.markdown("<p class='section-label'>Actual Hires by Month × L3 Org</p>", unsafe_allow_html=True)
        a_piv_l3 = fa.pivot_table(index=A_L3, columns="Month", aggfunc="size", fill_value=0)
        a_piv_l3 = a_piv_l3.reindex(sorted(a_piv_l3.columns), axis=1)
        a_piv_l3["Total"] = a_piv_l3.sum(axis=1)
        a_piv_l3 = a_piv_l3.sort_values("Total", ascending=False)
        a_grand_l3 = a_piv_l3.sum(); a_grand_l3.name = "Grand Total"
        a_piv_l3 = pd.concat([a_piv_l3, a_grand_l3.to_frame().T])
        a_piv_l3.index.name = "L3 Org"
        st.dataframe(a_piv_l3, use_container_width=True)

        st.markdown("---")

        # ── Table 2: Month × Job Level ────────────────────────────────────────
        st.markdown("<p class='section-label'>Actual Hires by Month × Job Level</p>", unsafe_allow_html=True)
        a_piv_jl = fa.pivot_table(index=A_JOB_LEVEL, columns="Month", aggfunc="size", fill_value=0)
        a_piv_jl = a_piv_jl.reindex(sorted(a_piv_jl.columns), axis=1)
        a_piv_jl["Total"] = a_piv_jl.sum(axis=1)
        a_piv_jl = a_piv_jl.sort_values("Total", ascending=False)
        a_grand_jl = a_piv_jl.sum(); a_grand_jl.name = "Grand Total"
        a_piv_jl = pd.concat([a_piv_jl, a_grand_jl.to_frame().T])
        a_piv_jl.index.name = "Job Level"
        st.dataframe(a_piv_jl, use_container_width=True)

        st.markdown("---")

        # ── Table 3: L3 Org × Job Level ───────────────────────────────────────
        st.markdown("<p class='section-label'>Actual Hires by L3 Org × Job Level</p>", unsafe_allow_html=True)
        a_piv_l3jl = fa.pivot_table(index=A_L3, columns=A_JOB_LEVEL, aggfunc="size", fill_value=0)
        a_piv_l3jl["Total"] = a_piv_l3jl.sum(axis=1)
        a_piv_l3jl = a_piv_l3jl.sort_values("Total", ascending=False)
        a_grand_l3jl = a_piv_l3jl.sum(); a_grand_l3jl.name = "Grand Total"
        a_piv_l3jl = pd.concat([a_piv_l3jl, a_grand_l3jl.to_frame().T])
        a_piv_l3jl.index.name = "L3 Org"
        st.dataframe(a_piv_l3jl, use_container_width=True)

        st.markdown("---")

        # ── Table 4: Hire Type & Band summary ─────────────────────────────────
        st.markdown("<p class='section-label'>Actual Hires by Hire Type & Band</p>", unsafe_allow_html=True)
        atc1, atc2 = st.columns(2)
        with atc1:
            aht_tbl = simple_summary(fa.replace("-", pd.NA).dropna(subset=[A_HIRE_TYPE]), A_HIRE_TYPE)
            aht_tbl.columns = ["Hire Type", "Count"]
            st.dataframe(aht_tbl, use_container_width=True, hide_index=True)
        with atc2:
            aband_tbl = simple_summary(fa.replace("-", pd.NA).dropna(subset=[A_BAND]), A_BAND)
            aband_tbl.columns = ["Band", "Count"]
            st.dataframe(aband_tbl, use_container_width=True, hide_index=True)

        st.markdown("---")

        # ── Detail table ──────────────────────────────────────────────────────
        st.markdown("<p class='section-label'>Actual Hires Detail</p>", unsafe_allow_html=True)
        st.caption(f"Sorted: L3 Org → Job Level → Hire Date  ·  {len(fa):,} rows")

        actual_col_map = {
            A_NAME:         "Employee Name",
            A_HIRE_DATE:    "Hire Date",
            A_HIRE_TYPE:    "Hire Type",
            A_BAND:         "Band",
            A_MANAGER:      "Manager",
            A_L2:           "Org Level 2",
            A_L3:           "Org Level 3",
            A_L4:           "Org Level 4",
            A_JOB_LEVEL:    "Job Level",
            A_JOB_TITLE:    "Job Title",
            A_JOB_FAMILY:   "Job Family",
            A_JOB_FAM_GRP:  "Job Family Group",
            A_CAREER_TRACK: "Career Track",
            A_COMMUNITY:    "Community",
            A_TECH_COMM:    "Tech Community",
            A_SITE:         "Business Site",
            A_COUNTRY:      "Country",
            A_FISCAL_MONTH: "Fiscal Month",
            A_FISCAL_YEAR:  "Fiscal Year",
        }
        acols = [c for c in actual_col_map if c in fa.columns]
        actual_tbl = (
            fa[acols].rename(columns=actual_col_map)
            .sort_values(["Org Level 3", "Job Level", "Hire Date"])
            .reset_index(drop=True)
        )

        def actual_row_style(row):
            if pd.notna(as_of_date) and pd.notna(row.get("Hire Date")) and row["Hire Date"] >= as_of_date:
                return ["background-color:#FFF3E0; color:#1A1A2E"] * len(row)
            return [""] * len(row)

        st.dataframe(
            actual_tbl.style.apply(actual_row_style, axis=1),
            use_container_width=True, height=480,
        )
        st.download_button(
            "Download Actual Hires YTD (CSV)",
            actual_tbl.to_csv(index=False).encode("utf-8"),
            "actual_hires_ytd_filtered.csv", "text/csv",
        )

# ══════════════════════════════════════════════════════════════════════════════
# REPORT GENERATION HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def _heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1C, 0x2B, 0x3A)
    run.font.name = "Arial"
    return p

def _para(doc, text, bold=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.name  = "Arial"
    run.font.bold  = bold
    run.font.color.rgb = RGBColor(0x1C, 0x2B, 0x3A)
    return p

def _df_to_table(doc, df, col_widths=None):
    df = df.reset_index(drop=False)
    tbl = doc.add_table(rows=1, cols=len(df.columns))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        run = hdr[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(hdr[i], "1C2B3A")
    for _, row in df.iterrows():
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            run.font.name = "Arial"
    if col_widths:
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = Inches(col_widths[i])
    return tbl

def _metric_table(doc, metrics: list):
    """metrics = list of (label, value) tuples — renders as a single-row card table."""
    tbl = doc.add_table(rows=2, cols=len(metrics))
    tbl.style = "Table Grid"
    for i, (label, value) in enumerate(metrics):
        top = tbl.rows[0].cells[i]
        top.text = str(label)
        r = top.paragraphs[0].runs[0]
        r.font.size = Pt(8); r.font.bold = True; r.font.name = "Arial"
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(top, "0077C5")
        bot = tbl.rows[1].cells[i]
        bot.text = str(value)
        r2 = bot.paragraphs[0].runs[0]
        r2.font.size = Pt(14); r2.font.bold = True; r2.font.name = "Arial"
        r2.font.color.rgb = RGBColor(0x1C, 0x2B, 0x3A)
        bot.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return tbl

def _rag_label(val, thresholds):
    """Return 🟢 / 🟡 / 🔴 based on (green_min, amber_min)."""
    g, a = thresholds
    if val >= g: return "🟢"
    if val >= a: return "🟡"
    return "🔴"

def build_report(
    doc_type: str,          # "Weekly" | "Monthly"
    period_label: str,
    stale_days: int,
    df_reqs,
    df_hires,
    df_actual,
    today: date,
) -> bytes:
    doc = Document()

    # ── Page margins (narrow) ────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ════════════════════════════════════════════════════════════════════════
    # COVER BLOCK
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("Talent Acquisition Report", 0).runs[0].font.color.rgb = RGBColor(0x1C, 0x2B, 0x3A)
    p = doc.add_paragraph()
    p.add_run(f"{doc_type} Report  ·  {period_label}").font.size = Pt(11)
    p.add_run(f"\nGenerated: {today.strftime('%d %b %Y')}").font.size = Pt(9)
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # COMPUTE KEY NUMBERS
    # ════════════════════════════════════════════════════════════════════════
    today_ts = pd.Timestamp(today)

    # Actual hires
    total_actual = len(df_actual) if df_actual is not None else 0
    actual_started = int((df_actual[A_HIRE_DATE] < today_ts).sum()) if df_actual is not None else 0
    actual_upcoming = total_actual - actual_started

    # Expected hires (plan proxy)
    expected_past = 0
    if df_hires is not None:
        expected_past = int((df_hires[H_START_DATE] < today_ts).sum())

    pct_plan = (actual_started / expected_past * 100) if expected_past > 0 else None

    # Pipeline health
    total_reqs   = len(df_reqs) if df_reqs is not None else 0
    stale_reqs   = int((df_reqs[R_DAYS_OPEN] >= stale_days).sum()) if df_reqs is not None else 0
    avg_days_open = float(df_reqs[R_DAYS_OPEN].mean()) if df_reqs is not None else 0
    total_open   = int(df_reqs[R_REMAINING].sum()) if df_reqs is not None else 0

    # Hiring velocity — month-over-month
    velocity_note = ""
    if df_actual is not None and len(df_actual) > 0:
        mo_counts = (
            df_actual.groupby("Month").size().reset_index(name="n").sort_values("Month")
        )
        if len(mo_counts) >= 2:
            last_n  = int(mo_counts.iloc[-1]["n"])
            prev_n  = int(mo_counts.iloc[-2]["n"])
            delta   = last_n - prev_n
            pct_chg = (delta / prev_n * 100) if prev_n > 0 else 0
            direction = "up" if delta >= 0 else "down"
            velocity_note = (
                f"{mo_counts.iloc[-1]['Month']}: {last_n} hires "
                f"({'+' if delta >= 0 else ''}{delta} / {pct_chg:+.0f}% vs prior month)"
            )
        elif len(mo_counts) == 1:
            last_n = int(mo_counts.iloc[-1]["n"])
            velocity_note = f"{mo_counts.iloc[-1]['Month']}: {last_n} hires (first month on record)"
            direction = "steady"
            delta = 0

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 1 — EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "Executive Summary", 1)

    # Build narrative
    plan_sentence = (
        f"Against the {expected_past} expected hires with start dates to date, "
        f"{actual_started} have been confirmed — representing "
        f"{pct_plan:.0f}% of plan."
        if pct_plan is not None
        else f"{actual_started} candidates have started to date, with {actual_upcoming} "
             f"more scheduled to start in the coming weeks."
    )
    pipeline_sentence = (
        f"The pipeline currently carries {total_reqs:,} open requisitions "
        f"({total_open:,} remaining openings), with an average of {avg_days_open:.0f} days open. "
        f"{stale_reqs} req{'s' if stale_reqs != 1 else ''} "
        f"{'have' if stale_reqs != 1 else 'has'} exceeded {stale_days} days "
        f"and {'require' if stale_reqs != 1 else 'requires'} immediate attention."
        if df_reqs is not None
        else "Pipeline data is not currently loaded."
    )
    velocity_sentence = (
        f"Hiring velocity: {velocity_note}."
        if velocity_note
        else "Insufficient data to calculate month-over-month hiring velocity."
    )

    _para(doc, f"{plan_sentence} {pipeline_sentence} {velocity_sentence}", size=10)
    doc.add_paragraph()

    # ── Headline metrics card ────────────────────────────────────────────────
    plan_display = f"{pct_plan:.0f}% of plan" if pct_plan is not None else "Plan TBD"
    _metric_table(doc, [
        ("Actual Hires YTD",      f"{total_actual:,}"),
        ("Already Started",       f"{actual_started:,}"),
        ("vs. Expected (to date)", plan_display),
        ("Open Reqs",             f"{total_reqs:,}"),
        ("Stale Reqs (60+ days)", f"{stale_reqs:,}"),
    ])
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 2 — WINS
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "Wins", 1)

    # Orgs at/above expected pace
    if df_actual is not None and df_hires is not None:
        act_by_l2 = df_actual.groupby(A_L2).size().reset_index(name="Actual")
        exp_by_l2 = (
            df_hires[df_hires[H_START_DATE] < today_ts]
            .groupby(H_L2).size().reset_index(name="Expected")
        ) if H_L2 in df_hires.columns else pd.DataFrame(columns=[H_L2, "Expected"])
        exp_by_l2 = exp_by_l2.rename(columns={H_L2: A_L2})
        pace = act_by_l2.merge(exp_by_l2, on=A_L2, how="left").fillna(0)
        pace["Expected"] = pace["Expected"].astype(int)
        pace["vs Plan"] = pace["Actual"] - pace["Expected"]
        ahead = pace[pace["vs Plan"] >= 0].sort_values("vs Plan", ascending=False)
        behind = pace[pace["vs Plan"] < 0].sort_values("vs Plan")
    else:
        ahead = pd.DataFrame()
        behind = pd.DataFrame()

    # MoM improvement by org
    mom_wins = pd.DataFrame()
    if df_actual is not None and len(df_actual) > 0:
        months_sorted = sorted(df_actual["Month"].unique())
        if len(months_sorted) >= 2:
            last_mo, prev_mo = months_sorted[-1], months_sorted[-2]
            last_counts = df_actual[df_actual["Month"] == last_mo].groupby(A_L3).size().reset_index(name="This Month")
            prev_counts = df_actual[df_actual["Month"] == prev_mo].groupby(A_L3).size().reset_index(name="Prior Month")
            mom = last_counts.merge(prev_counts, on=A_L3, how="outer").fillna(0)
            mom["This Month"]  = mom["This Month"].astype(int)
            mom["Prior Month"] = mom["Prior Month"].astype(int)
            mom["Change"]      = mom["This Month"] - mom["Prior Month"]
            mom_wins = mom[mom["Change"] > 0].sort_values("Change", ascending=False).rename(columns={A_L3: "L3 Org"})

    _para(doc, "Orgs at or above hiring pace (Actual ≥ Expected):", bold=True)
    if not ahead.empty:
        win_tbl = ahead[["Org Level 2", "Actual", "Expected", "vs Plan"]].rename(columns={"Org Level 2": "L2 Org"})
        _df_to_table(doc, win_tbl.head(8))
    else:
        _para(doc, "Data not available — upload both Expected and Actual Hires files.")
    doc.add_paragraph()

    _para(doc, f"Month-over-month improvement (L3 Org, {prev_mo if len(months_sorted) >= 2 else '—'} → {last_mo if len(months_sorted) >= 2 else '—'}):", bold=True)
    if not mom_wins.empty:
        _df_to_table(doc, mom_wins[["L3 Org", "Prior Month", "This Month", "Change"]].head(8))
    else:
        _para(doc, "Insufficient monthly data for MoM comparison (need 2+ months of actuals).")
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 3 — HOTSPOTS
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, "Hotspots", 1)

    # Stale reqs
    _para(doc, f"Open Requisitions — Stale (≥ {stale_days} days open):", bold=True)
    if df_reqs is not None:
        stale = df_reqs[df_reqs[R_DAYS_OPEN] >= stale_days].copy()
        if not stale.empty:
            stale_display = stale[[R_L2, R_L3, R_JOB_TITLE, R_JOB_LEVEL, R_DAYS_OPEN, R_REMAINING, R_RECRUITER]].copy()
            stale_display.columns = ["L2 Org", "L3 Org", "Job Title", "Level", "Days Open", "Remaining", "Recruiter"]
            stale_display = stale_display.sort_values("Days Open", ascending=False).reset_index(drop=True)
            _df_to_table(doc, stale_display.head(15))
        else:
            _para(doc, f"✅ No open reqs have exceeded {stale_days} days. Pipeline is healthy.")
    else:
        _para(doc, "Open Reqs file not loaded.")
    doc.add_paragraph()

    # Orgs behind pace
    _para(doc, "Orgs behind hiring pace (Actual < Expected):", bold=True)
    if not behind.empty:
        behind_tbl = behind[["Org Level 2", "Actual", "Expected", "vs Plan"]].rename(columns={"Org Level 2": "L2 Org"})
        _df_to_table(doc, behind_tbl.head(8))
    else:
        _para(doc, "All orgs are at or above expected pace — or Expected Hires file not loaded.")
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 4 — SUPPORTING DATA  (Monthly only / full depth)
    # ════════════════════════════════════════════════════════════════════════
    if doc_type == "Monthly":
        _heading(doc, "Supporting Data", 1)

        # Actual hires by month
        _para(doc, "Actual Hires by Month", bold=True)
        if df_actual is not None:
            mo_tbl = (
                df_actual.groupby(["Month", "Month Label"]).size()
                .reset_index(name="Hires").sort_values("Month")
                [["Month Label", "Hires"]]
            )
            mo_tbl.loc[len(mo_tbl)] = ["Total", mo_tbl["Hires"].sum()]
            _df_to_table(doc, mo_tbl.rename(columns={"Month Label": "Month"}))
        doc.add_paragraph()

        # Pipeline by L2 org
        _para(doc, "Open Requisitions by L2 Org", bold=True)
        if df_reqs is not None:
            l2_summary = (
                df_reqs.groupby(R_L2)
                .agg(Reqs=(R_PIPELINE_ID, "count"), Remaining=(R_REMAINING, "sum"), Avg_Days=(R_DAYS_OPEN, "mean"))
                .reset_index()
                .rename(columns={R_L2: "L2 Org", "Avg_Days": "Avg Days Open"})
                .sort_values("Remaining", ascending=False)
            )
            l2_summary["Avg Days Open"] = l2_summary["Avg Days Open"].round(0).astype(int)
            total_row = pd.DataFrame([{"L2 Org": "Total", "Reqs": l2_summary["Reqs"].sum(),
                                        "Remaining": l2_summary["Remaining"].sum(),
                                        "Avg Days Open": round(df_reqs[R_DAYS_OPEN].mean())}])
            _df_to_table(doc, pd.concat([l2_summary, total_row], ignore_index=True))
        doc.add_paragraph()

        # Actual hires by L3 org × job level
        _para(doc, "Actual Hires by L3 Org × Job Level", bold=True)
        if df_actual is not None:
            l3jl = df_actual.pivot_table(index=A_L3, columns=A_JOB_LEVEL, aggfunc="size", fill_value=0)
            l3jl["Total"] = l3jl.sum(axis=1)
            l3jl = l3jl.sort_values("Total", ascending=False)
            grand = l3jl.sum(); grand.name = "Grand Total"
            l3jl = pd.concat([l3jl, grand.to_frame().T])
            l3jl.index.name = "L3 Org"
            _df_to_table(doc, l3jl.reset_index())
        doc.add_paragraph()

        # Hire type breakdown
        _para(doc, "Actual Hires by Hire Type", bold=True)
        if df_actual is not None:
            ht_tbl = df_actual.groupby(A_HIRE_TYPE).size().reset_index(name="Count").sort_values("Count", ascending=False)
            ht_tbl.loc[len(ht_tbl)] = ["Total", ht_tbl["Count"].sum()]
            _df_to_table(doc, ht_tbl.rename(columns={A_HIRE_TYPE: "Hire Type"}))
        doc.add_paragraph()

    # ── Footer ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    _para(doc, f"Confidential · Talent Acquisition · Generated {today.strftime('%d %b %Y')} · Data as uploaded to Intuit Hiring Dashboard", size=8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# TAB 4 — REPORTS
# ─────────────────────────────────────────────────────────────────────────────
with tab4:
    st.markdown("### TA Director Report Generator")
    st.caption("Generate a Word document report you can share with your SVP. All data comes from the files already loaded in the dashboard.")

    any_data = df_reqs is not None or df_hires is not None or df_actual_hires is not None
    if not any_data:
        st.warning("Upload at least one data file in the sidebar to generate a report.")
    else:
        st.markdown("<div class='filter-card'>", unsafe_allow_html=True)

        rc1, rc2, rc3 = st.columns(3)
        with rc1:
            report_type = st.radio("Report depth", ["Weekly", "Monthly"],
                                   horizontal=True, key="rpt_type",
                                   help="Weekly = Exec summary + wins + hotspots (1 page). Monthly = adds full supporting data tables.")
        with rc2:
            period_label = st.text_input("Period label (appears on cover)",
                                          value=f"w/e {date.today().strftime('%d %b %Y')}" if True else "",
                                          placeholder="e.g. Week ending 28 Mar 2026",
                                          key="rpt_period")
        with rc3:
            stale_threshold = st.number_input("Stale req threshold (days)",
                                               min_value=1, max_value=365, value=60, step=5,
                                               key="rpt_stale")

        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("#### What will be in this report")
        preview_cols = st.columns(3)
        with preview_cols[0]:
            st.markdown("""
**📊 Executive Summary**
- Narrative paragraph with key numbers
- Hires YTD vs. expected (plan proxy)
- Pipeline health snapshot
- Hiring velocity (MoM)
""")
        with preview_cols[1]:
            st.markdown("""
**🏆 Wins**
- Orgs at or above hiring pace
- Month-over-month improvement by L3 org
""")
        with preview_cols[2]:
            depth = "Full supporting tables included" if report_type == "Monthly" else "Headline metrics only"
            st.markdown(f"""
**🔥 Hotspots**
- Reqs open ≥ {stale_threshold} days
- Orgs behind expected hiring pace

**📎 Supporting Data** ({'Monthly only' if report_type == 'Weekly' else '✅ included'})
- {depth}
""")

        st.markdown("---")

        if st.button(f"⬇️  Generate {report_type} Report", type="primary", key="rpt_generate"):
            with st.spinner("Building your report…"):
                try:
                    doc_bytes = build_report(
                        doc_type      = report_type,
                        period_label  = period_label or f"{report_type} — {date.today().strftime('%d %b %Y')}",
                        stale_days    = stale_threshold,
                        df_reqs       = df_reqs,
                        df_hires      = df_hires,
                        df_actual     = df_actual_hires,
                        today         = date.today(),
                    )
                    filename = f"TA_{report_type}_Report_{date.today().strftime('%Y%m%d')}.docx"
                    st.success(f"✅ Report ready — click below to download.")
                    st.download_button(
                        label    = f"📄 Download {report_type} Report (.docx)",
                        data     = doc_bytes,
                        file_name= filename,
                        mime     = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key      = "rpt_download",
                    )
                except Exception as e:
                    st.error(f"Report generation failed: {e}")
                    st.exception(e)
